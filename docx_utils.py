from docx import Document


def extract_docx_paragraphs(path):
    doc = Document(path)
    return [p.text for p in doc.paragraphs]


def rebuild_docx(original_path, translated_paragraphs, output_path):
    original = Document(original_path)
    new_doc = Document()

    for i, para in enumerate(original.paragraphs):
        new_para = new_doc.add_paragraph()
        new_para.style = para.style
        new_para.add_run(translated_paragraphs[i])

    new_doc.save(output_path)
